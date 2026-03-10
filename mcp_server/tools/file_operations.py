import os
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, Tuple, List
from enum import Enum
from mcp.server.fastmcp import Context
from .tool_base import ToolBase, ToolResult, OperationConfig, register_tool, ConfirmModel, extract_path_from_blackboard
from .security_sandbox import SecurityChecker, create_default_security_checker

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class FileOperationEnum(str, Enum):
    CREATE = "create"
    READ = "read"
    WRITE = "write"
    DELETE = "delete"
    MOVE = "move"
    COPY = "copy"
    LIST = "list"
    SEARCH = "search"
    CHECK_PERMISSION = "check_permission"
    READ_WRITE = "read_write"


class ReadWriteMode(str, Enum):
    R_PLUS = "r+"
    W_PLUS = "w+"
    A_PLUS = "a+"


@register_tool("file_operations")
class FileOperationsTool(ToolBase):
    """文件系统操作工具
    
    支持多种文件和文件夹操作，包括创建、读取、写入、删除、移动、复制、列出、搜索和读写。
    """
    
    TOOL_NAME = "file_operations"
    
    OPERATION_CONFIG = {
        'create': OperationConfig(
            description='创建文件/文件夹',
            required_params=['path'],
            optional_params=[],
            is_dangerous=False
        ),
        'read': OperationConfig(
            description='读取文件内容',
            required_params=['path'],
            optional_params=[],
            is_dangerous=False
        ),
        'write': OperationConfig(
            description='写入文件内容',
            required_params=['path', 'content'],
            optional_params=['overwrite'],
            is_dangerous=False
        ),
        'delete': OperationConfig(
            description='删除文件/文件夹',
            required_params=['path'],
            optional_params=[],
            is_dangerous=True
        ),
        'move': OperationConfig(
            description='移动文件/文件夹',
            required_params=['path', 'destination'],
            optional_params=[],
            is_dangerous=True
        ),
        'copy': OperationConfig(
            description='复制文件/文件夹',
            required_params=['path', 'destination'],
            optional_params=[],
            is_dangerous=False
        ),
        'list': OperationConfig(
            description='列出目录内容',
            required_params=['path'],
            optional_params=[],
            is_dangerous=False
        ),
        'search': OperationConfig(
            description='搜索文件',
            required_params=['path', 'content'],
            optional_params=[],
            is_dangerous=False
        ),
        'check_permission': OperationConfig(
            description='检查文件/文件夹权限',
            required_params=['path'],
            optional_params=[],
            is_dangerous=False
        ),
        'read_write': OperationConfig(
            description='读写文件内容',
            required_params=['path', 'content'],
            optional_params=['mode'],
            is_dangerous=False
        )
    }
    
    def __init__(self, security_checker: Optional[SecurityChecker] = None):
        super().__init__()
        self.security_checker = security_checker or create_default_security_checker()
    
    @classmethod
    def validate_parameters(cls, operation: str, **kwargs) -> Tuple[Dict[str, Any], Optional[str]]:
        """验证并调整参数"""
        params, config_error = super().validate_parameters(operation, **kwargs)
        
        if config_error:
            return params, config_error
        
        mode = kwargs.get('mode')
        if operation == "read_write" and mode:
            if mode not in [ReadWriteMode.R_PLUS, ReadWriteMode.W_PLUS, ReadWriteMode.A_PLUS]:
                config_error = f"不支持的读写模式: {mode}，支持的模式: r+、w+、a+"
        
        return params, config_error
    
    @staticmethod
    def process_path_static(path: str) -> str:
        """静态方法：处理文件路径，支持桌面路径、相对路径和绝对路径"""
        path = os.path.expanduser(path)
        desktop_path = str(Path.home() / "Desktop")
        
        if path == "桌面" or path == "desktop":
            return desktop_path
        elif path.startswith("桌面/") or path.startswith("desktop/"):
            return str(Path(desktop_path) / path.split("/", 1)[1])
        elif path.startswith("桌面\\") or path.startswith("desktop\\"):
            return str(Path(desktop_path) / path.split("\\", 1)[1])
        elif not os.path.isabs(path) and not path.startswith("./") and not path.startswith(".\\"):
            return str(Path(desktop_path) / path)
        else:
            return path
    
    def _process_path(self, path: str) -> str:
        """处理文件路径，支持桌面路径、相对路径和绝对路径"""
        return self.process_path_static(path)
    
    def _check_file_type(self, file_path: str) -> Tuple[str, bool]:
        """检查文件类型并判断是否支持
        
        Returns:
            (file_type, is_supported) - file_type: 'docx' 或 'text'，is_supported: 是否支持该文件类型
        """
        if file_path.lower().endswith('.docx'):
            return 'docx', DOCX_AVAILABLE
        return 'text', True
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """读取 Word 文件内容
        
        Returns:
            包含文件内容的字典，格式: {"content": str, "error": Optional[str]}
        """
        try:
            doc = Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            return {"content": '\n'.join(content), "error": None}
        except Exception as e:
            return {"content": "", "error": str(e)}
    
    def _write_docx(self, file_path: str, content: str, overwrite: bool) -> Dict[str, Any]:
        """写入 Word 文件内容
        
        Args:
            file_path: 文件路径
            content: 要写入的内容
            overwrite: 是否覆盖（True）或追加（False）
        
        Returns:
            包含操作结果的字典，格式: {"success": bool, "error": Optional[str]}
        """
        try:
            if overwrite or not os.path.exists(file_path):
                # 覆盖模式或文件不存在，创建新文档
                doc = Document()
            else:
                # 追加模式，打开现有文档
                doc = Document(file_path)
            
            # 将内容按段落分割
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            doc.save(file_path)
            return {"success": True, "error": None}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def _check_security(self, processed_path: str, operation: str, destination: str = None, ctx: Context = None) -> Tuple[bool, Optional[Dict[str, Any]]]:
        """安全检查
        
        Returns:
            (is_safe, error_result) - 安全返回(True, None)，不安全返回(False, error_dict)
        """
        if not self.security_checker:
            return True, None
        
        is_path_safe, path_error = self.security_checker.check_path(processed_path)
        if not is_path_safe:
            return False, ToolResult.error(path_error).build()
        
        if destination:
            dest_path = self._process_path(destination)
            is_dest_safe, dest_error = self.security_checker.check_path(dest_path)
            if not is_dest_safe:
                return False, ToolResult.error(dest_error).build()
        
        is_op_safe, op_error = self.security_checker.check_operation(operation)
        if not is_op_safe:
            return False, ToolResult.error(op_error).build()
        
        if self.security_checker.is_dangerous_operation(operation) and ctx:
            dangerous_message = f"检测到危险操作 '{operation}'，要操作的路径: {processed_path}，是否确认执行？"
            result = await ctx.elicit(message=dangerous_message, schema=ConfirmModel)
            if result.action != "accept" or not getattr(result.data, "confirmed", False):
                return False, ToolResult.error("用户取消执行").build()
        
        return True, None
    
    async def execute(self, ctx: Optional[Context] = None, **kwargs) -> Dict[str, Any]:
        """执行文件操作"""
        operation = kwargs.get('operation')
        path = kwargs.get('path')
        content = kwargs.get('content')
        destination = kwargs.get('destination')
        overwrite = kwargs.get('overwrite', True)
        mode = kwargs.get('mode')
        
        processed_path = self._process_path(path)
        
        # 强制安全检查，不受确认权限影响
        is_safe, error_result = await self._check_security(processed_path, operation, destination, ctx)
        
        if not is_safe:
            return error_result
        
        if operation == "create":
            return self._create(processed_path)
        elif operation == "write":
            return self._write(processed_path, content, overwrite)
        elif operation == "read":
            return self._read(processed_path)
        elif operation == "delete":
            if not await self._confirm_with_permission(
                ctx,
                f"确认删除文件/文件夹\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}",
                **kwargs
            ):
                return ToolResult.error("用户取消删除").build()
            
            return self._delete(processed_path)
        elif operation == "move":
            processed_destination = self._process_path(destination)
            if not await self._confirm_with_permission(
                ctx,
                f"确认移动文件/文件夹\n📄 源文件: {os.path.basename(processed_path)}\n📍 源路径: {processed_path}\n📄 目标文件: {os.path.basename(processed_destination)}\n📍 目标路径: {processed_destination}",
                **kwargs
            ):
                return ToolResult.error("用户取消移动").build()
            
            return self._move(processed_path, destination)
        elif operation == "copy":
            return self._copy(processed_path, destination)
        elif operation == "list":
            return self._list(processed_path)
        elif operation == "search":
            return self._search(processed_path, content)
        elif operation == "check_permission":
            return self._check_permission(processed_path)
        elif operation == "read_write":
            return self._read_write(processed_path, content, mode)
        else:
            return ToolResult.error(f"不支持的操作: {operation}").build()
    
    def _create(self, processed_path: str) -> Dict[str, Any]:
        """创建文件/文件夹"""
        if processed_path.endswith("/") or "." not in os.path.basename(processed_path):
            if os.path.exists(processed_path):
                return (ToolResult.success("文件夹已存在")
                    .with_path(processed_path)
                    .with_message(f"📁 文件夹已存在\n📍 路径: {processed_path}")
                    .with_blackboard(processed_path)
                    .build())
            os.makedirs(processed_path, exist_ok=True)
            return (ToolResult.success("文件夹创建成功")
                .with_path(processed_path)
                .with_message(f"✅ 文件夹创建成功\n📁 文件夹: {os.path.basename(processed_path)}\n📍 路径: {processed_path}")
                .with_blackboard(processed_path)
                .build())
        else:
            if os.path.exists(processed_path):
                return (ToolResult.success("文件已存在")
                    .with_path(processed_path)
                    .with_message(f"📄 文件已存在\n📍 路径: {processed_path}")
                    .with_blackboard(processed_path)
                    .build())
            dir_path = os.path.dirname(processed_path)
            if dir_path:
                os.makedirs(dir_path, exist_ok=True)
            with open(processed_path, 'w', encoding='utf-8') as f:
                f.write("")
            return (ToolResult.success("文件创建成功")
                .with_path(processed_path)
                .with_message(f"✅ 文件创建成功\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}")
                .with_blackboard(processed_path)
                .build())
    
    def _write(self, processed_path: str, content: str, overwrite: bool) -> Dict[str, Any]:
        """写入文件"""
        dir_path = os.path.dirname(processed_path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)
        
        # 检查文件类型
        file_type, is_supported = self._check_file_type(processed_path)
        
        if file_type == 'docx' and is_supported:
            # 使用统一的 Word 写入方法
            result = self._write_docx(processed_path, content or "", overwrite)
            if not result["success"]:
                return ToolResult.error(f"写入Word文件失败: {result['error']}").build()
            
            content_length = len(content or "")
            return (ToolResult.success("Word文件写入成功")
                .with_path(processed_path)
                .with_extra("write_mode", "追加" if not overwrite else "覆盖")
                .with_extra("content_length", content_length)
                .with_extra("file_type", "docx")
                .with_message(f"✅ Word文件写入成功\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}\n📝 写入模式: {'追加' if not overwrite else '覆盖'}\n📊 写入长度: {content_length} 字符")
                .with_blackboard(processed_path)
                .build())
        elif file_type == 'docx' and not is_supported:
            return ToolResult.error("缺少必要的库，请安装 python-docx").build()
        
        # 普通文本文件写入
        mode = 'a' if not overwrite else 'w'
        content_length = len(content or "")
        with open(processed_path, mode, encoding='utf-8') as f:
            f.write(content or "")
        
        return (ToolResult.success("文件写入成功")
            .with_path(processed_path)
            .with_extra("write_mode", "追加" if not overwrite else "覆盖")
            .with_extra("content_length", content_length)
            .with_extra("file_type", "text")
            .with_message(f"✅ 文件写入成功\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}\n📝 写入模式: {'追加' if not overwrite else '覆盖'}\n📊 写入长度: {content_length} 字符")
            .with_blackboard(processed_path)
            .build())
    
    def _read(self, processed_path: str) -> Dict[str, Any]:
        """读取文件"""
        # 使用通用函数处理路径（支持从列表字符串中提取第一个路径）
        try:
            processed_path = extract_path_from_blackboard(processed_path)
        except ValueError as e:
            return ToolResult.error(f"路径解析失败: {str(e)}").build()
        
        # 检查文件类型
        file_type, is_supported = self._check_file_type(processed_path)
        
        if file_type == 'docx' and is_supported:
            # 使用统一的 Word 读取方法
            result = self._read_docx(processed_path)
            if result["error"]:
                return ToolResult.error(f"读取Word文件失败: {result['error']}").build()
            
            content = result["content"]
            content_preview = content[:200] + ("..." if len(content) > 200 else "")
            return (ToolResult.success(content)
                .with_path(processed_path)
                .with_extra("content_length", len(content))
                .with_extra("file_type", "docx")
                .with_message(f"📄 Word文件读取成功\n📍 路径: {processed_path}\n📊 文件长度: {len(content)} 字符\n\n预览:\n{content_preview}")
                .with_blackboard(content)
                .build())
        elif file_type == 'docx' and not is_supported:
            return ToolResult.error("缺少必要的库，请安装 python-docx").build()
        
        # 普通文本文件读取
        with open(processed_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        content_preview = content[:200] + ("..." if len(content) > 200 else "")
        return (ToolResult.success(content)
            .with_path(processed_path)
            .with_extra("content_length", len(content))
            .with_extra("file_type", "text")
            .with_message(f"📄 文件读取成功\n📍 路径: {processed_path}\n📊 文件长度: {len(content)} 字符\n\n预览:\n{content_preview}")
            .with_blackboard(content)
            .build())
    
    def _delete(self, processed_path: str) -> Dict[str, Any]:
        """删除文件/文件夹"""
        is_dir = os.path.isdir(processed_path)
        item_type = "文件夹" if is_dir else "文件"
        
        if is_dir:
            shutil.rmtree(processed_path)
        else:
            os.remove(processed_path)
        
        return (ToolResult.success("删除成功")
            .with_path(processed_path)
            .with_message(f"✅ 删除成功\n{'📁' if is_dir else '📄'} {item_type}: {os.path.basename(processed_path)}")
            .with_blackboard(processed_path)
            .build())
    
    def _move(self, processed_path: str, destination: str) -> Dict[str, Any]:
        """移动文件/文件夹"""
        if not destination:
            return ToolResult.config_error("move操作需要destination参数").build()
        
        dest_path = self._process_path(destination)
        os.makedirs(os.path.dirname(dest_path), exist_ok=True)
        shutil.move(processed_path, dest_path)
        
        return (ToolResult.success("移动成功")
            .with_path(dest_path)
            .with_message(f"✅ 移动成功\n📄 项目: {os.path.basename(processed_path)}\n📍 目标路径: {dest_path}")
            .with_blackboard(dest_path)
            .build())
    
    def _copy(self, processed_path: str, destination: str) -> Dict[str, Any]:
        """复制文件/文件夹"""
        if not destination:
            return ToolResult.config_error("copy操作需要destination参数").build()
        
        dest_path = self._process_path(destination)
        os.makedirs(os.path.dirname(dest_path), exist_ok=True)
        
        is_dir = os.path.isdir(processed_path)
        if is_dir:
            shutil.copytree(processed_path, dest_path, dirs_exist_ok=True)
        else:
            shutil.copy2(processed_path, dest_path)
        
        return (ToolResult.success("复制成功")
            .with_path(dest_path)
            .with_message(f"✅ 复制成功\n{'📁' if is_dir else '📄'} {'文件夹' if is_dir else '文件'}: {os.path.basename(processed_path)}\n📍 目标路径: {dest_path}")
            .with_blackboard(dest_path)
            .build())
    
    def _list(self, processed_path: str) -> Dict[str, Any]:
        """列出目录内容"""
        items = os.listdir(processed_path)
        files = [item for item in items if os.path.isfile(os.path.join(processed_path, item))]
        folders = [item for item in items if os.path.isdir(os.path.join(processed_path, item))]
        
        formatted_items = []
        if folders:
            formatted_items.append("📁 文件夹:")
            for folder in folders[:10]:
                formatted_items.append(f"  - {folder}")
            if len(folders) > 10:
                formatted_items.append(f"  ... 等{len(folders) - 10}个文件夹")
        
        if files:
            formatted_items.append("\n📄 文件:")
            for file in files[:10]:
                formatted_items.append(f"  - {file}")
            if len(files) > 10:
                formatted_items.append(f"  ... 等{len(files) - 10}个文件")
        
        message = f"📋 目录内容\n📍 路径: {processed_path}\n📊 总计: {len(folders)}个文件夹, {len(files)}个文件\n\n" + "\n".join(formatted_items)
        
        return (ToolResult.success(items)
            .with_path(processed_path)
            .with_extra("folders", len(folders))
            .with_extra("files", len(files))
            .with_message(message)
            .with_blackboard(items)
            .build())
    
    def _search(self, processed_path: str, content: str) -> Dict[str, Any]:
        """搜索文件"""
        matches = []
        for root, dirs, files in os.walk(processed_path):
            depth = root[len(processed_path):].count(os.sep)
            if depth > 10:
                dirs[:] = []
                continue
            
            for file in files:
                if content and content.lower() in file.lower():
                    matches.append(os.path.join(root, file))
        
        formatted_matches = []
        for match in matches[:10]:
            formatted_matches.append(f"  - {os.path.basename(match)}")
        if len(matches) > 10:
            formatted_matches.append(f"  ... 等{len(matches) - 10}个文件")
        
        message = f"🔍 搜索结果\n📍 搜索路径: {processed_path}\n🔤 搜索关键词: {content}\n📊 找到 {len(matches)} 个匹配文件\n\n" + "\n".join(formatted_matches)
        
        return (ToolResult.success(matches[:50])
            .with_path(processed_path)
            .with_extra("total_matches", len(matches))
            .with_message(message)
            .with_blackboard(matches[:50])
            .build())
    
    def _check_permission(self, processed_path: str) -> Dict[str, Any]:
        """检查文件权限"""
        if not os.path.exists(processed_path):
            permission_data = {
                "exists": False,
                "readable": False,
                "writable": False,
                "executable": False,
                "is_file": False,
                "is_dir": False
            }
            return (ToolResult.success(permission_data)
                .with_path(processed_path)
                .with_message(f"🔍 文件权限检查完成\n📍 路径: {processed_path}\n❌ 文件/文件夹不存在")
                .with_blackboard(permission_data)
                .build())
        
        is_file = os.path.isfile(processed_path)
        is_dir = os.path.isdir(processed_path)
        readable = os.access(processed_path, os.R_OK)
        writable = os.access(processed_path, os.W_OK)
        executable = os.access(processed_path, os.X_OK)
        
        permission_status = []
        permission_status.append("✅ 可读" if readable else "❌ 不可读")
        permission_status.append("✅ 可写" if writable else "❌ 不可写")
        permission_status.append("✅ 可执行" if executable else "❌ 不可执行")
        
        permission_data = {
            "exists": True,
            "readable": readable,
            "writable": writable,
            "executable": executable,
            "is_file": is_file,
            "is_dir": is_dir
        }
        
        message = f"🔍 文件权限检查完成\n{'📄' if is_file else '📁'} {'文件' if is_file else '文件夹'}: {os.path.basename(processed_path)}\n📍 路径: {processed_path}\n\n权限状态:\n" + "\n".join(permission_status)
        
        return (ToolResult.success(permission_data)
            .with_path(processed_path)
            .with_message(message)
            .with_blackboard(permission_data)
            .build())
    
    def _read_write(self, processed_path: str, content: str, mode: Optional[ReadWriteMode]) -> Dict[str, Any]:
        """读写文件"""
        if not mode:
            mode = ReadWriteMode.W_PLUS
        
        # 检查文件类型
        file_type, is_supported = self._check_file_type(processed_path)
        
        if file_type == 'docx' and is_supported:
            # 使用 python-docx 处理 Word 文件
            return self._read_write_docx(processed_path, content, mode)
        elif file_type == 'docx' and not is_supported:
            return ToolResult.error("缺少必要的库，请安装 python-docx").build()
        
        # 普通文本文件处理
        if mode == ReadWriteMode.R_PLUS and not os.path.exists(processed_path):
            return ToolResult.error(f"r+模式需要文件存在，但文件 '{processed_path}' 不存在").build()
        
        dir_path = os.path.dirname(processed_path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)
        
        original_content = ""
        if os.path.exists(processed_path):
            try:
                with open(processed_path, 'r', encoding='utf-8') as f:
                    original_content = f.read()
            except Exception as e:
                return ToolResult.error(f"读取文件失败: {str(e)}").build()
        
        try:
            with open(processed_path, mode.value, encoding='utf-8') as f:
                if mode == ReadWriteMode.R_PLUS:
                    f.seek(0)
                    f.write(content or "")
                elif mode == ReadWriteMode.W_PLUS:
                    f.write(content or "")
                elif mode == ReadWriteMode.A_PLUS:
                    f.write(content or "")
            
            content_length = len(content or "")
            mode_desc = {
                ReadWriteMode.R_PLUS: '读写模式（文件必须存在）',
                ReadWriteMode.W_PLUS: '读写模式（覆盖文件）',
                ReadWriteMode.A_PLUS: '读写模式（追加内容）'
            }.get(mode, '未知模式')
            
            # 根据模式决定返回的 result_blackboard
            if mode == ReadWriteMode.R_PLUS:
                # R_PLUS 模式下，返回文件内容更合理
                try:
                    with open(processed_path, 'r', encoding='utf-8') as f:
                        updated_content = f.read()
                    blackboard_value = updated_content
                except Exception:
                    # 如果读取失败，返回文件路径
                    blackboard_value = processed_path
            else:
                # W_PLUS 和 A_PLUS 模式下，返回文件路径
                blackboard_value = processed_path
            
            return (ToolResult.success("文件读写成功")
                .with_path(processed_path)
                .with_extra("mode", mode_desc)
                .with_extra("content_length", content_length)
                .with_extra("file_type", "text")
                .with_message(f"✅ 文件读写成功\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}\n📝 模式: {mode_desc}\n📊 写入长度: {content_length} 字符")
                .with_blackboard(blackboard_value)
                .build())
        except Exception as e:
            return ToolResult.error(f"文件读写失败: {str(e)}").build()
    
    def _read_write_docx(self, processed_path: str, content: str, mode: ReadWriteMode) -> Dict[str, Any]:
        """读写 Word 文件"""
        dir_path = os.path.dirname(processed_path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)
        
        try:
            if mode == ReadWriteMode.R_PLUS:
                # 读写模式（文件必须存在）
                if not os.path.exists(processed_path):
                    return ToolResult.error(f"r+模式需要文件存在，但文件 '{processed_path}' 不存在").build()
                
                # 写入新内容（覆盖）
                result = self._write_docx(processed_path, content or "", True)
                if not result["success"]:
                    return ToolResult.error(f"Word文件写入失败: {result['error']}").build()
                
                # 读取更新后的内容
                result = self._read_docx(processed_path)
                if result["error"]:
                    return ToolResult.error(f"Word文件读取失败: {result['error']}").build()
                
                updated_content = result["content"]
                content_length = len(content or "")
                mode_desc = '读写模式（文件必须存在）'
                
                return (ToolResult.success("Word文件读写成功")
                    .with_path(processed_path)
                    .with_extra("mode", mode_desc)
                    .with_extra("content_length", content_length)
                    .with_extra("file_type", "docx")
                    .with_message(f"✅ Word文件读写成功\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}\n📝 模式: {mode_desc}\n📊 写入长度: {content_length} 字符")
                    .with_blackboard(updated_content)
                    .build())
            
            elif mode == ReadWriteMode.W_PLUS:
                # 读写模式（覆盖文件）
                result = self._write_docx(processed_path, content or "", True)
                if not result["success"]:
                    return ToolResult.error(f"Word文件写入失败: {result['error']}").build()
                
                content_length = len(content or "")
                mode_desc = '读写模式（覆盖文件）'
                
                return (ToolResult.success("Word文件读写成功")
                    .with_path(processed_path)
                    .with_extra("mode", mode_desc)
                    .with_extra("content_length", content_length)
                    .with_extra("file_type", "docx")
                    .with_message(f"✅ Word文件读写成功\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}\n📝 模式: {mode_desc}\n📊 写入长度: {content_length} 字符")
                    .with_blackboard(processed_path)
                    .build())
            
            elif mode == ReadWriteMode.A_PLUS:
                # 读写模式（追加内容）
                result = self._write_docx(processed_path, content or "", False)
                if not result["success"]:
                    return ToolResult.error(f"Word文件写入失败: {result['error']}").build()
                
                content_length = len(content or "")
                mode_desc = '读写模式（追加内容）'
                
                return (ToolResult.success("Word文件读写成功")
                    .with_path(processed_path)
                    .with_extra("mode", mode_desc)
                    .with_extra("content_length", content_length)
                    .with_extra("file_type", "docx")
                    .with_message(f"✅ Word文件读写成功\n📄 文件: {os.path.basename(processed_path)}\n📍 路径: {processed_path}\n📝 模式: {mode_desc}\n📊 写入长度: {content_length} 字符")
                    .with_blackboard(processed_path)
                    .build())
            
            else:
                return ToolResult.error(f"不支持的模式: {mode}").build()
            
        except Exception as e:
            return ToolResult.error(f"Word文件读写失败: {str(e)}").build()


def register_file_operations_tools(mcp, security_checker=None, output_callback=None):
    """注册文件系统操作工具到MCP服务器"""
    tool = FileOperationsTool(security_checker)
    
    @mcp.tool()
    async def file_operations(
        operation: FileOperationEnum,
        path: str,
        content: Optional[str] = None,
        destination: Optional[str] = None,
        overwrite: bool = True,
        mode: Optional[ReadWriteMode] = None,
        execution_mode: Optional[str] = None,
        ctx: Optional[Context] = None
    ) -> Dict[str, Any]:
        """文件系统操作工具
        
        支持多种文件和文件夹操作，包括创建、读取、写入、删除、移动、复制、列出、搜索和读写。
        
        Args:
            operation: 操作类型，可选值: create, read, write, delete, move, copy, list, search, check_permission, read_write
            path: 文件/文件夹路径
        
        write 操作参数:
            content: 写入内容（必需）
            overwrite: 是否覆盖（默认True覆盖，False追加）
        
        move/copy 操作参数:
            destination: 目标路径（必需）
        
        search 操作参数:
            content: 搜索关键词（必需）
        
        read_write 操作参数:
            content: 写入内容（必需）
            mode: 读写模式（默认w+）
                - "r+": 读写模式（文件必须存在）
                - "w+": 读写模式（覆盖文件）
                - "a+": 读写模式（追加内容）
        
        Returns:
            执行结果字典
        """
        return await tool.safe_execute(
            ctx=ctx,
            operation=operation,
            path=path,
            content=content,
            destination=destination,
            overwrite=overwrite,
            mode=mode,
            execution_mode=execution_mode
        )
    
    return file_operations
